import os
import math
import random
import pandas as pd

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import defaultdict


class GenerateIndividualTasks:

    def __init__(self, config):
        self.students_file = config["students_file"]
        self.questions_file = config["task_questions_file"]
        self.output_file = config["output_excel"]
        self.output_word_file = config["output_word"]

        self.max_semester = int(config["max_semester"])
        self.personalized_questions = bool(int(config["personalized_questions"]))
        self.add_extra_blank_page = bool(int(config["add_extra_blank_page"]))
        self.num_students = int(config["num_students"])
        self.questions_per_topic = config["questions_per_topic"]
        self.text = config["text"]

    def generate(self):
        try:
            if self.personalized_questions and os.path.exists(self.students_file):
                students = self.read_students(self.students_file)
            else:
                students = [f"Student {i + 1}" for i in range(self.num_students)]
            students.sort()

            questions_by_topic = self.read_questions_from_excel(self.questions_file, self.max_semester)
            student_variants = self.generate_variants(students, questions_by_topic, self.questions_per_topic)

            self.save_to_excel(student_variants, self.output_file)
            self.create_word_document(student_variants, self.output_word_file)

            print(f"Documents have been created successfully: {self.output_file}, {self.output_word_file}")
        except Exception as e:
            print(f"Error: {e}")

    @staticmethod
    def read_students(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return [
                ' '.join(line.strip().split())
                for line in file.readlines()
                if line.strip()
            ]

    @staticmethod
    def read_questions_from_excel(file_path, max_semester):
        df = pd.read_excel(
            file_path,
            header=None,
            names=["Question", "Topic", "Semester", "Type"],
            usecols=range(4)
        )

        filtered_df = df[df["Semester"] <= max_semester]
        questions_by_topic = defaultdict(list)

        for _, row in filtered_df.iterrows():
            questions_by_topic[row["Topic"]].append((row["Question"], row["Type"]))

        return questions_by_topic

    @staticmethod
    def generate_variants(students, all_questions_by_topic, questions_per_topic):

        student_variants = defaultdict(list)
        question_count = defaultdict(int)

        max_repeats_by_topic = {
            topic: math.ceil(len(students) / len(all_questions_by_topic[topic])) + 2
            for topic in all_questions_by_topic
        }

        for student in students:
            assigned_questions = set()

            for topic, num_questions in questions_per_topic.items():
                questions_by_topic = set()

                while len(questions_by_topic) < num_questions:
                    question, question_type = random.choice(all_questions_by_topic[topic])

                    if question_count[question] >= max_repeats_by_topic[topic]:
                        all_questions_by_topic[topic] = [
                            q for q in all_questions_by_topic[topic]
                            if q[0] != question
                        ]

                    if not all_questions_by_topic[topic]:
                        raise ValueError(f"Not enough questions in topic '{topic}'")

                    questions_by_topic.add((question, question_type))
                    question_count[question] += 1

                assigned_questions.update(questions_by_topic)

            student_variants[student] = list(assigned_questions)

        return student_variants

    @staticmethod
    def save_to_excel(student_variants, output_path):
        df = pd.DataFrame.from_dict(
            student_variants,
            orient='index'
        ).reset_index()

        df.columns = ['Student'] + [
            f'Question {i+1}'
            for i in range(len(df.columns) - 1)
        ]

        df.to_excel(output_path, index=False)

        print(f"Variants saved to {output_path}")

    @staticmethod
    def add_page_break(document):
        page_break = OxmlElement('w:br')
        page_break.set(qn('w:type'), 'page')
        document.add_paragraph()._element.append(page_break)

    @staticmethod
    def set_cell_border(cell, **kwargs):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        for border_name in ["top", "left", "bottom", "right"]:
            if border_name in kwargs:
                border = OxmlElement(f"w:{border_name}")
                for key, value in kwargs[border_name].items():
                    border.set(qn(f"w:{key}"), str(value))
                tcPr.append(border)

    @staticmethod
    def set_document_spacing(document):
        for paragraph in document.paragraphs:
            paragraph.paragraph_format.space_after = Pt(4)

    def create_word_document(self, students_questions, output_path):
        document = Document()
        for section in document.sections:
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.text = self.text["task_header_title"]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.text = self.text["task_footer_title"]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for section in document.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        for count, (student, questions) in enumerate(students_questions.items(), start=1):
            document.add_paragraph()
            title_paragraph = document.add_paragraph()
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title_paragraph.add_run(self.text["task_ticket_title"] + f" №{count}")
            run.bold = True
            run.font.size = Pt(16)
            fio_group_paragraph = document.add_paragraph()

            if self.personalized_questions:
                fio_group_paragraph.add_run(f"{student}").bold = True
            else:
                fio_group_paragraph.add_run(self.text["task_name_title"]).bold = True

            num_questions = len(questions)
            table = document.add_table(
                rows=2,
                cols=num_questions + 1
            )

            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            for i in range(num_questions):
                cell = table.cell(0, i)
                cell.text = str(i + 1)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            sum_cell = table.cell(0, num_questions)
            sum_cell.text = self.text["task_sum_symbol"]
            sum_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for row in table.rows:
                for cell in row.cells:
                    self.set_cell_border(
                        cell,
                        top={"sz": 6, "val": "single", "color": "000000"},
                        bottom={"sz": 6, "val": "single", "color": "000000"},
                        left={"sz": 6, "val": "single", "color": "000000"},
                        right={"sz": 6, "val": "single", "color": "000000"}
                    )

            explanation_paragraph = document.add_paragraph()
            explanation_run = explanation_paragraph.add_run(self.text["task_info"])
            explanation_run.font.size = Pt(9)
            explanation_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            theory_paragraph = document.add_paragraph()
            theory_run = theory_paragraph.add_run(self.text["task_theory_section"])
            theory_run.bold = True
            theory_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            theory_questions = [q for q in questions if q[1] == "Theory"]

            for i, (question, _) in enumerate(theory_questions, start=1):
                document.add_paragraph(
                    f"{i}. {question}"
                ).paragraph_format.line_spacing = Pt(10)

            practical_questions = [q for q in questions if q[1] == "Practical"]
            if len(practical_questions) > 0:
                practical_paragraph = document.add_paragraph()
                practical_run = practical_paragraph.add_run(self.text["task_practical_section"])
                practical_run.bold = True
                practical_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for i, (question, _) in enumerate(practical_questions, start=1):
                    document.add_paragraph(
                        f"{i+len(theory_questions)}. {question}"
                    ).paragraph_format.line_spacing = Pt(10)
                    document.add_paragraph()

            self.add_page_break(document)

            if self.add_extra_blank_page:
                title_paragraph = document.add_paragraph()
                title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = title_paragraph.add_run(self.text["task_extra_form_title"] + f" №{count}")
                run.bold = True
                run.font.size = Pt(16)
                if self.personalized_questions:
                    fio_group_paragraph.add_run(f"{student}").bold = True
                else:
                    fio_group_paragraph.add_run(self.text["task_name_title"]).bold = True

                for _ in range(10):
                    document.add_paragraph()

                self.add_page_break(document)

        self.set_document_spacing(document)
        document.save(output_path)
